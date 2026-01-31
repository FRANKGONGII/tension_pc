import os

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.shared import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def print_doc(now_handle_data_id=-1):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    def set_table_border(table, top=False, bottom=False, left=False, right=False, thickness=10):
        """设置表格边框粗细
        """
        # 获取表格属性
        tbl = table._tbl
        tbl_pr = tbl.tblPr
        
        # 创建或获取表格边框设置
        tbl_borders = tbl_pr.find(qn('w:tblBorders'))
        if tbl_borders is None:
            tbl_borders = OxmlElement('w:tblBorders')
            tbl_pr.append(tbl_borders)
        
        # 设置上下边框
        if top:
            top_border = OxmlElement('w:top')
            top_border.set(qn('w:val'), 'single')
            top_border.set(qn('w:sz'), str(thickness))
            top_border.set(qn('w:space'), '0')
            top_border.set(qn('w:color'), 'auto')
            # 移除现有的top边框（如果存在）
            for e in tbl_borders.findall(qn('w:top')):
                tbl_borders.remove(e)
            tbl_borders.append(top_border)
            
        if bottom:
            bottom_border = OxmlElement('w:bottom')
            bottom_border.set(qn('w:val'), 'single')
            bottom_border.set(qn('w:sz'), str(thickness))
            bottom_border.set(qn('w:space'), '0')
            bottom_border.set(qn('w:color'), 'auto')
            # 移除现有的bottom边框（如果存在）
            for e in tbl_borders.findall(qn('w:bottom')):
                tbl_borders.remove(e)
            tbl_borders.append(bottom_border)
        
        # 设置左右边框
        if left:
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'single')
            left_border.set(qn('w:sz'), str(thickness))
            left_border.set(qn('w:space'), '0')
            left_border.set(qn('w:color'), 'auto')
            # 移除现有的left边框（如果存在）
            for e in tbl_borders.findall(qn('w:left')):
                tbl_borders.remove(e)
            tbl_borders.append(left_border)
        
        if right:
            right_border = OxmlElement('w:right')
            right_border.set(qn('w:val'), 'single')
            right_border.set(qn('w:sz'), str(thickness))
            right_border.set(qn('w:space'), '0')
            right_border.set(qn('w:color'), 'auto')
            # 移除现有的right边框（如果存在）
            for e in tbl_borders.findall(qn('w:right')):
                tbl_borders.remove(e)
            tbl_borders.append(right_border)

    # 查询数据库获取数据
    from utils.data_manager import DataManager
    detail = DataManager.queryById(now_handle_data_id)
    print(now_handle_data_id, detail)

    # 创建文档
    doc = Document()
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    doc.styles['Title'].font.name = 'SimSun'
    doc.styles['Title']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    doc.styles['Heading 1'].font.name = 'SimSun'
    doc.styles['Heading 1']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # 获取第一个节（默认文档只有一个节）
    section = doc.sections[0]

    # 设置页边距（单位：英寸），紧凑布局确保一页内
    section.top_margin = Inches(0.3)
    section.bottom_margin = Inches(0.3)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    footer = section.footer
    # 添加固定的页脚文本
    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 居中对齐
    run = paragraph.add_run('第 1 页')
    run.font.size = Pt(8)  # 设置字体大小



    # 添加公司名和标题（居中）
    def set_simsun_font(run, size=12, bold=False):
        """为指定的run设置宋体字体"""
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.size = Pt(size)
        run.bold = bold
        return run

    # 添加公司名和标题（居中）并设置宋体
    def add_centered_text_with_simsun(content, font_size=12, bold=False, style=None):
        paragraph = doc.add_paragraph(style=style)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(content)
        set_simsun_font(run, font_size, bold)
        run.font.color.rgb = RGBColor(0, 0, 0)  # 设置为黑色
        # 设置段前段后为0，行距为单倍
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0)
        paragraph_format.line_spacing = 1.0
        return paragraph

    # 添加中文标题（大号宋体，加粗）
    add_centered_text_with_simsun("江苏慧通管道设备股份有限公司", font_size=18, bold=True, style='Heading 1')

    # 添加英文公司名（正常宋体）
    add_centered_text_with_simsun("JiangShu Huitong Pipeline Equipment Co.Ltd", font_size=14)

    # 添加主标题（大号宋体，加粗）
    add_centered_text_with_simsun("恒力吊架性能试验记录", font_size=14, bold=True, style='Heading 1')

    # 添加英文副标题（正常宋体）
    add_centered_text_with_simsun("Constant Hanger Performance Test Record", font_size=12)

    # 添加基本信息表格
    table1 = doc.add_table(rows=3, cols=6)
    table1.style = 'Table Grid'


    # TODO：边框设置加粗
    # detail字段顺序见data_manager.py表结构
    # (id, test_time, user, 吊点代号, 出厂编号, 型号规格, 工作荷载, 位移方向, 总位移, 工作位移, 操作员, 检验员, 位移起始点值, 位移终止点值, 实测位移值, 超载试验值, 起止时间, 保持时间, 恒定度, 锁定位置，载荷偏差度, 测试结果)
    table1.rows[0].cells[0].text = "用户\nCustomer"
    table1.rows[0].cells[2].text = "规格型号\nDescription and type"
    table1.rows[0].cells[4].text = "试验日期\nTest Date"
    # 第二行内容 吊点代号 总位移 工作载荷
    table1.rows[1].cells[0].text = "吊点代号\nHanger Number"
    table1.rows[1].cells[2].text = "总位移\nTotal travel"
    table1.rows[1].cells[4].text = "工作载荷\nWorking Load"
    # 第三行内容 出场编号 位移方向 工作位移
    table1.rows[2].cells[0].text = "出厂编号\nSerial Number"
    table1.rows[2].cells[2].text = "位移方向\nTravel direction"
    table1.rows[2].cells[4].text = "工作位移\nWorking travel"

    def safe_str(val):
        return "" if val is None else str(val)

    table1.rows[0].cells[1].text = safe_str(detail[2]) if detail else ""
    table1.rows[0].cells[3].text = safe_str(detail[5]) if detail else ""
    table1.rows[0].cells[5].text = safe_str(detail[1]) if detail else ""

    table1.rows[1].cells[1].text = safe_str(detail[3]) if detail else ""
    table1.rows[1].cells[3].text = safe_str(detail[8]) + "mm" if detail else ""
    table1.rows[1].cells[5].text = safe_str(detail[6]) + "N" if detail else ""

    table1.rows[2].cells[1].text = safe_str(detail[4]) if detail else ""
    table1.rows[2].cells[3].text = safe_str(detail[7]) if detail else ""
    table1.rows[2].cells[5].text = safe_str(detail[9]) + "mm" if detail else ""
    format_table_cells(table1, font_size=10)
    # 设置第一个表格的上边和左右边加粗
    set_table_border(table1, top=True, left=True, right=True)

    # 添加一个 1x1 的表格
    table = doc.add_table(rows=1, cols=1)
    table.style = 'Table Grid'
    # 设置这个表格的左右边加粗
    set_table_border(table, left=True, right=True)
    # 获取单元格
    cell = table.cell(0, 0)
    # 在单元格内添加图片
    paragraph = cell.paragraphs[0]  # 获取单元格内的段落
    picture = paragraph.add_run().add_picture('./resources/png.png', width=Inches(6))  # 图片约6英寸
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(4)
    paragraph_format.space_after = Pt(4)



    # 查询test_data用于Pmax/Pmin
    x_list, y_list = DataManager.queryTestDataByFormId(now_handle_data_id)
    def format_float(val):
        try:
            return f"{float(val):.3f}"
        except Exception:
            return safe_str(val)

    pmax = format_float(max(x_list)) if x_list else ""
    pmin = format_float(min(x_list)) if x_list else ""

    # 第一行单独为2x12表格，横向排列6个2x2小表格
    top_table = doc.add_table(rows=2, cols=12)
    top_table.style = 'Table Grid'
    # 操作员
    top_table.cell(0, 0).merge(top_table.cell(1, 1))
    top_table.cell(0, 0).text = "操作员\nOperator"
    # 刘云佳
    top_table.cell(0, 2).merge(top_table.cell(1, 3))
    top_table.cell(0, 2).text = safe_str(detail[10]) if detail else ""
    # 实测力值
    top_table.cell(0, 4).merge(top_table.cell(1, 5))
    top_table.cell(0, 4).text = "实测力值\nActual load"
    # Pmax/Pmin不合并
    top_table.cell(0, 6).text = "Pmax"
    top_table.cell(0, 7).text = "Pmin"
    top_table.cell(1, 6).text = pmax
    top_table.cell(1, 7).text = pmin
    # 检验员
    top_table.cell(0, 8).merge(top_table.cell(1, 9))
    top_table.cell(0, 8).text = "检验员"
    # 陈广春
    top_table.cell(0, 10).merge(top_table.cell(1, 11))
    top_table.cell(0, 10).text = safe_str(detail[11]) if detail else ""
    format_table_cells(top_table, font_size=10)
    # 设置这个表格的左右边加粗
    set_table_border(top_table, left=True, right=True)

    # 其余内容为一个2x6表格
    main_table = doc.add_table(rows=2, cols=6)
    main_table.style = 'Table Grid'
    main_table.cell(0, 0).text = "位移起始点值\nInitial point"
    main_table.cell(0, 1).text = safe_str(detail[12]) if detail else ""
    main_table.cell(0, 2).text = "位移终止点值\nFinishing point"
    main_table.cell(0, 3).text = safe_str(detail[13]) if detail else ""
    main_table.cell(0, 4).text = "实测位移值\nActual travel"
    main_table.cell(0, 5).text = safe_str(detail[14]) if detail else ""

    main_table.cell(1, 0).text = "超载实验值\nOverload\ntest data"
    main_table.cell(1, 1).text = safe_str(detail[15]) if detail else ""
    main_table.cell(1, 2).text = "超载起始-终止时间\ntime of\nstarting-finishing"
    main_table.cell(1, 3).text = safe_str(detail[16]) if detail else ""
    main_table.cell(1, 4).text = "超载实验保持时间\nDuration within\noverload test"
    main_table.cell(1, 5).text = safe_str(detail[17]) if detail else ""
    format_table_cells(main_table, font_size=10)
    # 设置这个表格的左右边加粗
    set_table_border(main_table, left=True, right=True)

    # 最后一个1行的表格，4个内容（自动适应页面宽度与上方表格一致）
    bottom_table = doc.add_table(rows=1, cols=8)
    bottom_table.style = 'Table Grid'
    # 不设置任何列宽，让表格完全自动适应页面宽度，与 main_table 一致
    bottom_table.cell(0, 0).text = "恒定度\nConstant rate"
    bottom_table.cell(0, 1).text = safe_str(detail[18]) if detail else ""
    bottom_table.cell(0, 2).text = "锁定位置\nLoad tolerance"
    bottom_table.cell(0, 3).text = safe_str(detail[19]) if detail else ""
    bottom_table.cell(0, 4).text = "荷载偏差度\nTest result"
    bottom_table.cell(0, 5).text = safe_str(detail[20]) if detail else ""
    bottom_table.cell(0, 6).text = "测试结果\nTest result"
    bottom_table.cell(0, 7).text = safe_str(detail[21]) if detail else ""
    format_table_cells(bottom_table, font_size=10)
    # 设置最后一个表格的下边和左右边加粗
    set_table_border(bottom_table, bottom=True, left=True, right=True)

    # 保存 Word 文件
    doc.save("恒力吊架性能试验记录" + str(now_handle_data_id) + ".docx")

    import win32com.client

    def print_word_file(path):
        word = win32com.client.Dispatch("Word.Application")
        word.ActivePrinter = "Canon iP1188 series"
        doc = word.Documents.Open(path)
        doc.PrintOut()  # 打印
        # ActivePrinter="EPSON4D76BB (L4160 Series)"
        doc.Close(False)
        word.Quit()
    print_word_file(os.getcwd() + "/恒力吊架性能试验记录" + str(now_handle_data_id) + ".docx")




def format_table_cells(table, font_size=9):
    for row in table.rows:
        for cell in row.cells:
            # 设置垂直居中
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for paragraph in cell.paragraphs:
                # 设置水平居中
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(font_size)


import pypandoc

def convert_docx_to_pdf(input_path, output_path):
    output = pypandoc.convert_file(input_path, 'pdf', outputfile=output_path)
    return output_path if output == "" else None


if __name__ == '__main__':
    # pdf_path = convert_docx_to_pdf("恒力吊架性能试验记录（无图）.docx", "output.pdf")
    # print("转换成功：" + pdf_path)
    print_doc()
